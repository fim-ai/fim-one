"""Conversation export -- MD, TXT, DOCX, PDF."""

from __future__ import annotations

import io
import json
import logging
import re
from datetime import datetime, timezone
from enum import Enum
from zoneinfo import ZoneInfo
from html.parser import HTMLParser
from typing import Any
from urllib.parse import quote

from fastapi import APIRouter, Depends, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from fim_one.db import get_session
from fim_one.web.auth import get_current_user
from fim_one.web.exceptions import AppError
from fim_one.web.export_fonts import (
    DOCX_LATIN_FONT,
    DOCX_MONO_FONT,
    PdfFonts,
    docx_east_asian_font,
    resolve_pdf_fonts,
)
from fim_one.db.models import Conversation, Message, User

router = APIRouter(prefix="/api/conversations", tags=["export"])
logger = logging.getLogger(__name__)


class ExportFormat(str, Enum):
    MD = "md"
    TXT = "txt"
    DOCX = "docx"
    PDF = "pdf"


class DetailLevel(str, Enum):
    FULL = "full"
    SUMMARY = "summary"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _sanitize_filename(title: str) -> str:
    """Replace characters unsafe for filenames with underscores."""
    safe = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", title)
    safe = safe.strip(". ")
    return safe or "conversation"


def _format_date(dt: datetime | None, tz_name: str | None = None) -> str:
    """Return a date string, optionally converted to the given timezone."""
    if dt is None:
        return ""
    if tz_name:
        try:
            tz = ZoneInfo(tz_name)
            # If datetime is naive (no tzinfo), assume UTC
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            dt = dt.astimezone(tz)
        except (KeyError, Exception):
            pass  # Fall back to raw datetime if timezone is invalid
    return dt.strftime("%Y-%m-%d %H:%M")


def _format_date_compact(dt: datetime | None, tz_name: str | None = None) -> str:
    """Return a compact date for filenames like 20260309."""
    if dt is None:
        return "export"
    if tz_name:
        try:
            tz = ZoneInfo(tz_name)
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            dt = dt.astimezone(tz)
        except (KeyError, Exception):
            pass
    return dt.strftime("%Y%m%d")


# ---------------------------------------------------------------------------
# I18n — export label translations
# ---------------------------------------------------------------------------

_I18N: dict[str, dict[str, str]] = {
    "en": {
        "mode": "Mode",
        "mode_standard": "Standard",
        "mode_planner": "Planner",
        "created": "Created",
        "total_tokens": "Total Tokens",
        "turn": "Turn",
        "user": "User",
        "assistant": "Assistant",
        "execution_details": "Execution Details",
        "iteration": "iteration",
        "iterations": "iterations",
        "iteration_label": "Iteration",
        "step": "Step",
        "tool": "Tool",
        "duration": "Duration",
        "reasoning": "Reasoning",
        "arguments": "Arguments",
        "result": "Result",
        "plan": "Plan",
        "round": "Round",
        "task": "Task",
        "dependencies": "Dependencies",
        "none": "none",
        "goal_achieved": "Goal Achieved",
        "goal_not_achieved": "Goal Not Achieved — No Steps Completed",
        "analysis": "Analysis",
        "confidence": "confidence",
        "step_result": "Step Result",
        "empty_conversation": "(empty conversation)",
        "untitled": "Untitled Conversation",
    },
    "zh": {
        "mode": "模式",
        "mode_standard": "标准",
        "mode_planner": "规划",
        "created": "创建时间",
        "total_tokens": "总 Token 数",
        "turn": "轮次",
        "user": "用户",
        "assistant": "助手",
        "execution_details": "执行详情",
        "iteration": "次迭代",
        "iterations": "次迭代",
        "iteration_label": "迭代",
        "step": "步骤",
        "tool": "工具",
        "duration": "耗时",
        "reasoning": "推理",
        "arguments": "参数",
        "result": "结果",
        "plan": "计划",
        "round": "轮",
        "task": "任务",
        "dependencies": "依赖",
        "none": "无",
        "goal_achieved": "目标已达成",
        "goal_not_achieved": "目标未达成",
        "analysis": "分析",
        "confidence": "置信度",
        "step_result": "步骤结果",
        "empty_conversation": "（空对话）",
        "untitled": "无标题对话",
    },
}


def _t(key: str, locale: str) -> str:
    """Get a translated string for the given locale."""
    strings = _I18N.get(locale, _I18N["en"])
    return strings.get(key, _I18N["en"].get(key, key))


def _mode_label(mode: str, locale: str = "en") -> str:
    return _t("mode_planner", locale) if mode == "dag" else _t("mode_standard", locale)


def _resolve_effective_mode(conv: Conversation, messages: list[Message]) -> str:
    """Resolve 'auto' mode to actual execution mode by inspecting message events.

    Conversations created via the auto-router may still have ``mode='auto'``
    in the DB (for records created before the auto_endpoint fix).  We detect
    the effective mode by looking for DAG-specific ``phase`` events in the
    assistant messages.
    """
    if conv.mode != "auto":
        return conv.mode
    # Check assistant messages for DAG-specific events (phase events)
    for msg in messages:
        if msg.role != "assistant":
            continue
        events = _extract_sse_events(msg)
        for ev in events:
            if ev.get("event") == "phase":
                return "dag"
    return "react"


def _pair_messages(messages: list[Message]) -> list[dict[str, Any]]:
    """Group messages into turns: each turn has a user message and an
    optional assistant message.  Messages are expected to be sorted by
    ``created_at``."""
    turns: list[dict[str, Any]] = []
    current_turn: dict[str, Any] | None = None

    for msg in messages:
        if msg.role == "user":
            current_turn = {"user": msg, "assistant": None}
            turns.append(current_turn)
        elif msg.role == "assistant":
            if current_turn is not None:
                current_turn["assistant"] = msg
            else:
                # Orphan assistant message -- create a turn without user msg
                turns.append({"user": None, "assistant": msg})

    return turns


# Broad Unicode emoji ranges — covers SMP pictograph blocks, BMP symbol
# blocks commonly rendered with emoji presentation, and joiners/selectors.
_EMOJI_RE = re.compile(
    "["
    # ---- Supplemental Multilingual Plane (all pictograph blocks) ----
    "\U0001F000-\U0001FBFF"  # Mahjong → Playing Cards → Enclosed Alphanum Sup →
    #                          Enclosed Ideographic Sup → Misc Symbols & Pictographs →
    #                          Emoticons → Ornamental Dingbats → Transport/Map →
    #                          Supplemental Symbols → Geometric Shapes Ext →
    #                          Symbols Extended-A → Legacy Computing
    # ---- Basic Multilingual Plane – emoji-capable blocks ----
    "\U00002139"             # ℹ information source
    "\U00002194-\U00002199"  # ↔↕↖↗↘↙ arrows
    "\U000021A9-\U000021AA"  # ↩↪ curved arrows
    "\U0000231A-\U000023FF"  # Misc Technical (⌚⏰⏳⏺ etc.)
    "\U000024C2"             # Ⓜ circled M
    "\U000025A0-\U000025FF"  # Geometric Shapes (▶◻◼◽◾ etc.)
    "\U00002600-\U000027BF"  # Misc Symbols + Dingbats
    "\U00002934-\U00002935"  # ⤴⤵ curved arrows
    "\U00002B05-\U00002B07"  # ⬅⬆⬇ arrows
    "\U00002B1B-\U00002B1C"  # ⬛⬜ large squares
    "\U00002B50"             # ⭐ star
    "\U00002B55"             # ⭕ circle
    "\U00003030"             # 〰 wavy dash
    "\U0000303D"             # 〽 part alternation mark
    "\U00003297"             # ㊗ congratulation
    "\U00003299"             # ㊙ secret
    # ---- Joiners, modifiers, selectors ----
    "\U0000200B-\U0000200D"  # zero-width space / non-joiner / joiner
    "\U0000FE00-\U0000FE0F"  # variation selectors
    "\U000020E3"             # combining enclosing keycap
    "\U000E0020-\U000E007F"  # tags (flag subdivisions)
    "]+",
)


def _strip_emoji(text: str) -> str:
    """Remove emoji characters that DOCX/PDF fonts cannot render."""
    return _EMOJI_RE.sub("", text)


# ---------------------------------------------------------------------------
# File attachment stripping
# ---------------------------------------------------------------------------

# Regex patterns for file content blocks and annotations injected by
# the upload pipeline.  These must be stripped before exporting user
# messages so that raw document text does not leak into exports.
_FILE_BLOCK_RE = re.compile(
    r"\n*--- File: .+? \(file_id: [^)]+\) ---[\s\S]*",
)
_ATTACHED_FILES_RE = re.compile(
    r"\n*\[Attached files[^\]]*\]",
)
_ATTACHED_IMAGES_RE = re.compile(
    r"\n*\[Attached images \([^)]*\): [^\]]*\]",
)


def _strip_attachments(content: str) -> str:
    """Remove file content blocks and attachment annotations from message content."""
    if not content:
        return content
    # Strip file content blocks (everything from first --- File: to end)
    content = _FILE_BLOCK_RE.sub("", content)
    # Strip [Attached files ...] annotations
    content = _ATTACHED_FILES_RE.sub("", content)
    # Strip [Attached images ...] annotations
    content = _ATTACHED_IMAGES_RE.sub("", content)
    return content.rstrip()


# ---------------------------------------------------------------------------
# SSE event parsing helpers
# ---------------------------------------------------------------------------


def _extract_sse_events(msg: Message) -> list[dict[str, Any]]:
    """Return the sse_events list from an assistant message's metadata."""
    meta = msg.metadata_ or {}
    return list(meta.get("sse_events", []))


def _extract_react_steps(events: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Extract completed tool-call steps from ReAct SSE events."""
    steps: list[dict[str, Any]] = []
    for ev in events:
        if ev.get("event") == "step":
            data = ev.get("data", {})
            if data.get("type") == "tool_call":
                steps.append(data)
    return steps


def _extract_done_event(events: list[dict[str, Any]]) -> dict[str, Any] | None:
    """Extract the done event from SSE events."""
    for ev in events:
        if ev.get("event") == "done":
            return dict(ev.get("data", {}))
    return None


def _extract_dag_plan(events: list[dict[str, Any]], target_round: int = 1) -> list[dict[str, Any]]:
    """Extract plan steps for a given DAG round."""
    for ev in events:
        if ev.get("event") == "phase":
            data = ev.get("data", {})
            if (
                data.get("name") == "planning"
                and data.get("status") == "done"
                and data.get("round", 1) == target_round
            ):
                return list(data.get("steps", []))
    return []


def _extract_dag_step_details(events: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    """Group step_progress events by step_id, collecting iterations and completion info."""
    steps: dict[str, dict[str, Any]] = {}
    for ev in events:
        if ev.get("event") == "step_progress":
            data = ev.get("data", {})
            sid = data.get("step_id", "")
            if sid not in steps:
                steps[sid] = {"iterations": [], "completed": None, "task": data.get("task", "")}

            event_type = data.get("event")
            if event_type == "iteration":
                steps[sid]["iterations"].append(data)
            elif event_type == "completed":
                steps[sid]["completed"] = data
            elif event_type == "started":
                steps[sid]["task"] = data.get("task", steps[sid]["task"])

    return steps


def _extract_dag_analysis(events: list[dict[str, Any]]) -> dict[str, Any] | None:
    """Extract the analysis phase result."""
    for ev in events:
        if ev.get("event") == "phase":
            data = ev.get("data", {})
            if data.get("name") == "analyzing" and data.get("status") == "done":
                return dict(data)
    return None


def _extract_dag_rounds(events: list[dict[str, Any]]) -> list[int]:
    """Determine which DAG rounds exist in the events."""
    rounds: set[int] = set()
    for ev in events:
        if ev.get("event") == "phase":
            data = ev.get("data", {})
            r = data.get("round")
            if r is not None:
                rounds.add(r)
    return sorted(rounds) if rounds else [1]


# ---------------------------------------------------------------------------
# Markdown renderer
# ---------------------------------------------------------------------------


def _render_md(conv: Conversation, messages: list[Message], detail: DetailLevel, locale: str = "en", tz_name: str | None = None) -> str:
    """Render a conversation as a Markdown document."""
    effective_mode = _resolve_effective_mode(conv, messages)
    lines: list[str] = []
    title = conv.title or _t("untitled", locale)

    # Header
    lines.append(f"# {title}")
    lines.append("")
    meta_parts = [
        f"**{_t('mode', locale)}:** {_mode_label(effective_mode, locale)}",
    ]
    meta_parts.append(f"**{_t('created', locale)}:** {_format_date(conv.created_at, tz_name)}")
    lines.append(" | ".join(meta_parts))

    if detail == DetailLevel.FULL and conv.total_tokens:
        lines.append(f"**{_t('total_tokens', locale)}:** {conv.total_tokens:,}")

    lines.append("")
    lines.append("---")
    lines.append("")

    # Turns
    turns = _pair_messages(messages)
    for idx, turn in enumerate(turns, 1):
        lines.append(f"## {_t('turn', locale)} {idx}")
        lines.append("")

        # User message
        user_msg: Message | None = turn.get("user")
        if user_msg:
            lines.append(f"**{_t('user', locale)}:**")
            lines.append("")
            lines.append(_strip_attachments(user_msg.content or ""))
            lines.append("")

        # Assistant message
        asst_msg: Message | None = turn.get("assistant")
        if asst_msg is None:
            lines.append("---")
            lines.append("")
            continue

        lines.append(f"**{_t('assistant', locale)}:**")
        lines.append("")

        if detail == DetailLevel.FULL:
            events = _extract_sse_events(asst_msg)
            done = _extract_done_event(events)

            if effective_mode == "dag":
                _render_md_dag_details(lines, events, done, locale)
            else:
                _render_md_react_details(lines, events, done, locale)

        # Final answer
        answer = asst_msg.content or ""
        if not answer and asst_msg.metadata_:
            answer = asst_msg.metadata_.get("answer", "")
        lines.append(answer)
        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def _render_md_react_details(
    lines: list[str],
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
    locale: str = "en",
) -> None:
    """Append ReAct execution details in Markdown."""
    steps = _extract_react_steps(events)
    if not steps:
        return

    iterations = done.get("iterations", len(steps)) if done else len(steps)
    elapsed = done.get("elapsed", 0) if done else 0

    iter_label = _t("iteration" if iterations == 1 else "iterations", locale)
    lines.append(
        f"> **{_t('execution_details', locale)}** ({iterations} {iter_label}, {elapsed:.1f}s)"
    )
    lines.append(">")
    lines.append(f"> | # | {_t('tool', locale)} | {_t('duration', locale)} |")
    lines.append("> |---|------|----------|")
    for i, step in enumerate(steps, 1):
        tool = step.get("tool_name", "unknown")
        dur = step.get("iter_elapsed", 0)
        lines.append(f"> | {i} | {tool} | {dur:.1f}s |")
    lines.append("")

    for i, step in enumerate(steps, 1):
        tool = step.get("tool_name", "unknown")
        dur = step.get("iter_elapsed", 0)
        lines.append(f"<details>")
        lines.append(f"<summary>{_t('step', locale)} {i}: {tool} ({dur:.1f}s)</summary>")
        lines.append("")

        reasoning = step.get("reasoning")
        if reasoning:
            lines.append(f"**{_t('reasoning', locale)}:** {reasoning}")
            lines.append("")

        args = step.get("tool_args")
        if args:
            lines.append(f"**{_t('arguments', locale)}:**")
            lines.append("```json")
            lines.append(json.dumps(args, indent=2, ensure_ascii=False))
            lines.append("```")
            lines.append("")

        observation = step.get("observation")
        if observation:
            lines.append(f"**{_t('result', locale)}:**")
            lines.append(str(observation))
            lines.append("")

        lines.append("</details>")
        lines.append("")


def _render_md_dag_details(
    lines: list[str],
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
    locale: str = "en",
) -> None:
    """Append DAG execution details in Markdown."""
    rounds = _extract_dag_rounds(events)

    for rnd in rounds:
        plan_steps = _extract_dag_plan(events, rnd)
        if plan_steps:
            lines.append(f"> **{_t('plan', locale)}** ({_t('round', locale)} {rnd})")
            lines.append(">")
            lines.append(f"> | {_t('step', locale)} | {_t('task', locale)} | {_t('dependencies', locale)} |")
            lines.append("> |------|------|-------------|")
            for ps in plan_steps:
                sid = ps.get("id", "?")
                task = ps.get("task", "")
                deps = ", ".join(ps.get("deps", [])) or "\u2014"
                lines.append(f"> | {sid} | {task} | {deps} |")
            lines.append("")

    step_details = _extract_dag_step_details(events)
    for sid, info in step_details.items():
        task = info.get("task", "")
        completed = info.get("completed")
        status = completed.get("status", "completed") if completed else "unknown"
        duration = completed.get("duration", 0) if completed else 0

        lines.append("<details>")
        lines.append(f"<summary>{sid}: {task} ({status}, {duration:.1f}s)</summary>")
        lines.append("")

        iterations = info.get("iterations", [])
        if iterations:
            lines.append(f"**{_t('iteration_label', locale)}:**")
            for it in iterations:
                it_num = it.get("iteration", "?")
                tool = it.get("tool_name", "unknown")
                dur = it.get("iter_elapsed", 0)
                reasoning = it.get("reasoning", "")
                observation = it.get("observation", "")
                lines.append(
                    f"{it_num}. **{tool}** ({dur:.1f}s) \u2014 {_t('reasoning', locale)}: {reasoning}"
                )
                if observation:
                    lines.append(f"   {_t('result', locale)}: {observation}")
            lines.append("")

        result = completed.get("result", "") if completed else ""
        if result:
            lines.append(f"**{_t('step_result', locale)}:** {result}")
            lines.append("")

        lines.append("</details>")
        lines.append("")

    analysis = _extract_dag_analysis(events)
    if analysis:
        achieved = analysis.get("achieved", False)
        confidence = analysis.get("confidence", 0)
        reasoning = analysis.get("reasoning", "")
        achieved_label = _t("goal_achieved", locale) if achieved else _t("goal_not_achieved", locale)
        lines.append(
            f"> **{_t('analysis', locale)}:** {achieved_label} ({confidence * 100:.0f}% {_t('confidence', locale)})"
        )
        if reasoning:
            lines.append(f"> {reasoning}")
        lines.append("")


# ---------------------------------------------------------------------------
# TXT renderer
# ---------------------------------------------------------------------------


def _strip_md(text: str) -> str:
    """Naively strip Markdown formatting for plain-text output."""
    # Remove HTML tags
    text = re.sub(r"<[^>]+>", "", text)
    # Remove bold/italic markers
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    # Remove heading markers
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Remove blockquote markers
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    # Remove code fences
    text = re.sub(r"```\w*\n?", "", text)
    # Remove horizontal rules
    text = re.sub(r"^---+\s*$", "=" * 60, text, flags=re.MULTILINE)
    return text


def _render_txt(conv: Conversation, messages: list[Message], detail: DetailLevel, locale: str = "en", tz_name: str | None = None) -> str:
    """Render a conversation as plain text."""
    effective_mode = _resolve_effective_mode(conv, messages)
    lines: list[str] = []
    title = conv.title or _t("untitled", locale)

    lines.append(title)
    lines.append("=" * len(title))
    lines.append("")
    lines.append(f"{_t('mode', locale)}: {_mode_label(effective_mode, locale)}")
    lines.append(f"{_t('created', locale)}: {_format_date(conv.created_at, tz_name)}")
    if detail == DetailLevel.FULL and conv.total_tokens:
        lines.append(f"{_t('total_tokens', locale)}: {conv.total_tokens:,}")
    lines.append("")
    lines.append("=" * 60)
    lines.append("")

    turns = _pair_messages(messages)
    for idx, turn in enumerate(turns, 1):
        lines.append(f"{_t('turn', locale)} {idx}")
        lines.append("-" * 40)
        lines.append("")

        user_msg: Message | None = turn.get("user")
        if user_msg:
            lines.append(f"[{_t('user', locale)}]")
            lines.append("")
            lines.append(_strip_attachments(user_msg.content or ""))
            lines.append("")

        asst_msg: Message | None = turn.get("assistant")
        if asst_msg is None:
            lines.append("=" * 60)
            lines.append("")
            continue

        lines.append(f"[{_t('assistant', locale)}]")
        lines.append("")

        if detail == DetailLevel.FULL:
            events = _extract_sse_events(asst_msg)
            done = _extract_done_event(events)

            if effective_mode == "dag":
                _render_txt_dag_details(lines, events, done, locale)
            else:
                _render_txt_react_details(lines, events, done, locale)

        answer = asst_msg.content or ""
        if not answer and asst_msg.metadata_:
            answer = asst_msg.metadata_.get("answer", "")
        lines.append(answer)
        lines.append("")
        lines.append("=" * 60)
        lines.append("")

    return "\n".join(lines)


def _render_txt_react_details(
    lines: list[str],
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
    locale: str = "en",
) -> None:
    """Append ReAct details as plain text."""
    steps = _extract_react_steps(events)
    if not steps:
        return

    iterations = done.get("iterations", len(steps)) if done else len(steps)
    elapsed = done.get("elapsed", 0) if done else 0

    iter_label = _t("iteration" if iterations == 1 else "iterations", locale)
    lines.append(
        f"  {_t('execution_details', locale)}: {iterations} {iter_label}, {elapsed:.1f}s"
    )
    lines.append("")

    for i, step in enumerate(steps, 1):
        tool = step.get("tool_name", "unknown")
        dur = step.get("iter_elapsed", 0)
        lines.append(f"  {_t('step', locale)} {i}: {tool} ({dur:.1f}s)")

        reasoning = step.get("reasoning")
        if reasoning:
            lines.append(f"    {_t('reasoning', locale)}: {reasoning}")

        args = step.get("tool_args")
        if args:
            lines.append(f"    {_t('arguments', locale)}: {json.dumps(args, ensure_ascii=False)}")

        observation = step.get("observation")
        if observation:
            lines.append(f"    {_t('result', locale)}: {observation}")

        lines.append("")


def _render_txt_dag_details(
    lines: list[str],
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
    locale: str = "en",
) -> None:
    """Append DAG details as plain text."""
    rounds = _extract_dag_rounds(events)

    for rnd in rounds:
        plan_steps = _extract_dag_plan(events, rnd)
        if plan_steps:
            lines.append(f"  {_t('plan', locale)} ({_t('round', locale)} {rnd}):")
            for ps in plan_steps:
                sid = ps.get("id", "?")
                task = ps.get("task", "")
                deps = ", ".join(ps.get("deps", [])) or _t("none", locale)
                lines.append(f"    {sid}: {task} [{_t('dependencies', locale)}: {deps}]")
            lines.append("")

    step_details = _extract_dag_step_details(events)
    for sid, info in step_details.items():
        task = info.get("task", "")
        completed = info.get("completed")
        status = completed.get("status", "completed") if completed else "unknown"
        duration = completed.get("duration", 0) if completed else 0

        lines.append(f"  {sid}: {task} ({status}, {duration:.1f}s)")

        for it in info.get("iterations", []):
            it_num = it.get("iteration", "?")
            tool = it.get("tool_name", "unknown")
            dur = it.get("iter_elapsed", 0)
            reasoning = it.get("reasoning", "")
            observation = it.get("observation", "")
            lines.append(f"    {_t('iteration_label', locale)} {it_num}: {tool} ({dur:.1f}s)")
            if reasoning:
                lines.append(f"      {_t('reasoning', locale)}: {reasoning}")
            if observation:
                lines.append(f"      {_t('result', locale)}: {observation}")

        result = completed.get("result", "") if completed else ""
        if result:
            lines.append(f"    {_t('step_result', locale)}: {result}")
        lines.append("")

    analysis = _extract_dag_analysis(events)
    if analysis:
        achieved = analysis.get("achieved", False)
        confidence = analysis.get("confidence", 0)
        reasoning = analysis.get("reasoning", "")
        label = _t("goal_achieved", locale) if achieved else _t("goal_not_achieved", locale)
        lines.append(f"  {_t('analysis', locale)}: {label} ({confidence * 100:.0f}% {_t('confidence', locale)})")
        if reasoning:
            lines.append(f"    {reasoning}")
        lines.append("")


# ---------------------------------------------------------------------------
# Shared typography constants
# ---------------------------------------------------------------------------

#: Scripts whose text has no inter-word spaces and therefore needs ReportLab's
#: character-level ``wordWrap="CJK"``. Latin-only paragraphs keep the default
#: word-level wrapping so English is not broken mid-word.
_CJK_RE = re.compile(
    "[ᄀ-ᇿ⺀-〿぀-ヿ㄰-㆏㐀-䶿"
    "一-鿿가-힯豈-﫿︰-﹏＀-￯]"
)


def _has_cjk(text: str) -> bool:
    return bool(_CJK_RE.search(text))


# Palette shared by the PDF and DOCX renderers (matches the portal's amber UI).
_INK = "#1F1B16"
_MUTED = "#6B5D4F"
_ACCENT = "#946B2D"
_RULE = "#C8B898"
_CODE_BG = "#F8F4ED"


# ---------------------------------------------------------------------------
# Markdown -> HTML helper (shared by DOCX and PDF renderers)
# ---------------------------------------------------------------------------


def _md_to_html(text: str) -> str:
    """Convert markdown to HTML using the markdown library."""
    import markdown  # type: ignore[import-untyped]

    return str(markdown.markdown(text, extensions=["fenced_code", "tables", "nl2br"]))


# ---------------------------------------------------------------------------
# DOCX styling
# ---------------------------------------------------------------------------

#: Paragraph style carrying the "User" / "Assistant" role labels. Kept out of
#: the Heading ladder so that headings written by the assistant have room
#: beneath it without inverting the hierarchy.
_DOCX_ROLE_STYLE = "FIM Role"

#: Heading level the assistant's own markdown headings start at. Levels 1-2 are
#: reserved for the document title and the turn separators.
_DOCX_CONTENT_HEADING_OFFSET = 2

#: ``(point size, colour)`` per Word Heading level, 1-6.
_DOCX_HEADING_SCALE: dict[int, tuple[float, str]] = {
    1: (20, _ACCENT),
    2: (15, _ACCENT),
    3: (13.5, _INK),
    4: (12, _INK),
    5: (11.5, _INK),
    6: (11, _INK),
}


#: Header-row fill for DOCX tables, matching the PDF renderer.
_TABLE_HEADER_BG = "F5EDDE"


def _shade_cell(cell: Any, hex_fill: str) -> None:
    """Apply a solid background fill to a table cell."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _hex_to_rgb(value: str) -> Any:
    from docx.shared import RGBColor

    return RGBColor.from_string(value.lstrip("#").upper())


def _apply_rpr_fonts(rpr: Any, latin: str, east_asian: str) -> None:
    """Stamp both the Latin and the East Asian font slot on an ``rPr``.

    Word keeps separate font slots per script.  python-docx writes only the
    Latin ones, so CJK text silently falls through to the theme font — a
    family that often has no bold face, which is why mixed Latin/CJK bold runs
    come out half-bold.
    """
    from docx.oxml.ns import qn

    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asian)
    rfonts.set(qn("w:cs"), latin)


def _set_docx_default_fonts(doc: Any, latin: str, east_asian: str) -> None:
    """Set the document-wide default fonts (``w:docDefaults``)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    element = doc.styles.element
    doc_defaults = element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        element.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    _apply_rpr_fonts(rpr, latin, east_asian)


def _configure_docx_styles(doc: Any, locale: str) -> None:
    """Install the export's typographic scale on a fresh document.

    python-docx starts from Word's default template: 11pt Calibri body, a
    Heading ladder whose sizes barely separate (16/13/12pt) and no East Asian
    font anywhere.  This replaces it with an explicit five-tier scale and
    pins both font slots on every style the exporter uses.
    """
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt

    east_asian = docx_east_asian_font(locale)

    _set_docx_default_fonts(doc, DOCX_LATIN_FONT, east_asian)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _hex_to_rgb(_INK)
    _apply_rpr_fonts(normal.element.get_or_add_rPr(), DOCX_LATIN_FONT, east_asian)
    fmt = normal.paragraph_format
    fmt.line_spacing = 1.4
    fmt.space_after = Pt(6)

    for level, (size, color) in _DOCX_HEADING_SCALE.items():
        style = doc.styles[f"Heading {level}"]
        style.font.size = Pt(size)
        style.font.bold = True
        # Word's built-in Heading 4 is italic, which reads as a slanted synthetic
        # face for CJK. Weight and size carry the hierarchy instead.
        style.font.italic = False
        style.font.color.rgb = _hex_to_rgb(color)
        _apply_rpr_fonts(style.element.get_or_add_rPr(), DOCX_LATIN_FONT, east_asian)
        style.paragraph_format.space_before = Pt(max(6.0, size * 0.8))
        style.paragraph_format.space_after = Pt(max(4.0, size * 0.35))

    try:
        role = doc.styles[_DOCX_ROLE_STYLE]
    except KeyError:
        role = doc.styles.add_style(_DOCX_ROLE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    role.base_style = doc.styles["Normal"]
    role.font.size = Pt(10)
    role.font.bold = True
    role.font.color.rgb = _hex_to_rgb(_ACCENT)
    _apply_rpr_fonts(role.element.get_or_add_rPr(), DOCX_LATIN_FONT, east_asian)
    role.paragraph_format.space_before = Pt(10)
    role.paragraph_format.space_after = Pt(2)


def _docx_content_heading_level(markdown_level: int) -> int:
    """Map an ``<h1>``-``<h6>`` from message content onto a Word heading level."""
    return min(markdown_level + _DOCX_CONTENT_HEADING_OFFSET, 6)


# ---------------------------------------------------------------------------
# DOCX markdown renderer
# ---------------------------------------------------------------------------


class _DocxMarkdownRenderer(HTMLParser):
    """Parse HTML (converted from Markdown) and emit python-docx elements."""

    def __init__(self, doc: Any, east_asian_font: str) -> None:
        super().__init__()
        self._doc = doc
        self._east_asian_font = east_asian_font
        self._paragraph: Any | None = None
        self._bold = False
        self._italic = False
        self._code_inline = False
        self._in_pre = False
        self._pre_text = ""
        self._heading_level = 0
        self._heading_text = ""
        self._in_blockquote = False
        self._list_style: str | None = None
        self._list_stack: list[str] = []
        self._in_table = False
        self._table_rows: list[list[str]] = []
        self._current_row: list[str] = []
        self._current_cell_text = ""
        self._in_td = False
        self._in_th = False
        self._href: str | None = None

    def _ensure_paragraph(self, style: str | None = None) -> Any:
        if self._paragraph is None:
            self._paragraph = self._doc.add_paragraph(style=style)
        return self._paragraph

    def _finish_paragraph(self) -> None:
        self._paragraph = None

    def _add_run(self, text: str) -> None:
        from docx.shared import Pt, RGBColor

        para = self._ensure_paragraph()
        run = para.add_run(text)
        if self._bold:
            run.bold = True
        if self._italic:
            run.italic = True
        if self._in_blockquote:
            # Colour, not italics: synthetic obliques look wrong on CJK glyphs.
            run.font.color.rgb = RGBColor.from_string(_MUTED.lstrip("#").upper())
        if self._code_inline:
            run.font.size = Pt(9.5)
            _apply_rpr_fonts(
                run._element.get_or_add_rPr(), DOCX_MONO_FONT, self._east_asian_font
            )
        if self._href:
            run.underline = True
            run.font.color.rgb = RGBColor(0x94, 0x6B, 0x2D)

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_dict = dict(attrs)

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._heading_level = int(tag[1])
            self._heading_text = ""
        elif tag == "p":
            style = None
            if self._list_stack:
                style = self._list_stack[-1]
            self._paragraph = self._doc.add_paragraph(style=style)
            if self._in_blockquote:
                from docx.shared import Inches

                self._paragraph.paragraph_format.left_indent = Inches(0.3)
        elif tag in ("strong", "b"):
            self._bold = True
        elif tag in ("em", "i"):
            self._italic = True
        elif tag == "code":
            if self._in_pre:
                pass  # handled by pre
            else:
                self._code_inline = True
        elif tag == "pre":
            self._in_pre = True
            self._pre_text = ""
        elif tag == "ul":
            self._list_stack.append("List Bullet")
        elif tag == "ol":
            self._list_stack.append("List Number")
        elif tag == "li":
            if self._list_stack:
                self._paragraph = self._doc.add_paragraph(style=self._list_stack[-1])
            else:
                self._paragraph = self._doc.add_paragraph(style="List Bullet")
        elif tag == "table":
            self._in_table = True
            self._table_rows = []
        elif tag == "tr":
            self._current_row = []
        elif tag in ("td", "th"):
            self._in_td = tag == "td"
            self._in_th = tag == "th"
            self._current_cell_text = ""
        elif tag == "blockquote":
            self._in_blockquote = True
        elif tag == "a":
            self._href = attrs_dict.get("href", "")
        elif tag == "hr":
            self._finish_paragraph()
            # Native DOCX horizontal rule via paragraph bottom border
            p = self._doc.add_paragraph()
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "C8B898")
            pBdr.append(bottom)
            pPr.append(pBdr)
            self._finish_paragraph()
        elif tag == "br":
            if self._paragraph:
                self._paragraph.add_run("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._doc.add_heading(
                self._heading_text,
                level=_docx_content_heading_level(self._heading_level),
            )
            self._heading_level = 0
            self._heading_text = ""
        elif tag == "p":
            self._finish_paragraph()
        elif tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False
        elif tag == "code":
            if not self._in_pre:
                self._code_inline = False
        elif tag == "pre":
            self._in_pre = False
            _add_monospace_paragraph(self._doc, self._pre_text, self._east_asian_font)
            self._pre_text = ""
        elif tag == "ul":
            if self._list_stack:
                self._list_stack.pop()
        elif tag == "ol":
            if self._list_stack:
                self._list_stack.pop()
        elif tag == "li":
            self._finish_paragraph()
        elif tag in ("td", "th"):
            self._current_row.append(self._current_cell_text.strip())
            self._in_td = False
            self._in_th = False
            self._current_cell_text = ""
        elif tag == "tr":
            self._table_rows.append(self._current_row)
            self._current_row = []
        elif tag == "table":
            self._in_table = False
            if self._table_rows:
                self._emit_table()
        elif tag == "blockquote":
            self._in_blockquote = False
        elif tag == "a":
            self._href = None

    def handle_charref(self, name: str) -> None:
        """Convert HTML numeric character references (&#123; / &#xAB;) to
        characters and route through ``handle_data``."""
        try:
            codepoint = int(name[1:], 16) if name.startswith(("x", "X")) else int(name)
            self.handle_data(chr(codepoint))
        except (ValueError, OverflowError):
            pass

    def handle_data(self, data: str) -> None:
        data = _strip_emoji(data)
        if self._heading_level:
            self._heading_text += data
        elif self._in_pre:
            self._pre_text += data
        elif self._in_td or self._in_th:
            self._current_cell_text += data
        elif self._in_table:
            pass  # ignore whitespace between table elements
        else:
            if data.strip() or self._paragraph:
                self._add_run(data)

    def _emit_table(self) -> None:
        from docx.shared import Pt

        if not self._table_rows:
            return
        n_cols = max(len(r) for r in self._table_rows) if self._table_rows else 1
        table = self._doc.add_table(rows=len(self._table_rows), cols=n_cols)
        table.style = "Table Grid"
        table.autofit = True

        # Bold header row
        for c_idx, cell_text in enumerate(self._table_rows[0] if self._table_rows else []):
            if c_idx < n_cols:
                cell = table.rows[0].cells[c_idx]
                cell.text = ""
                run = cell.paragraphs[0].add_run(cell_text)
                run.bold = True
                _shade_cell(cell, _TABLE_HEADER_BG)

        # Data rows
        for r_idx, row in enumerate(self._table_rows[1:], 1):
            for c_idx, cell_text in enumerate(row):
                if c_idx < n_cols:
                    table.rows[r_idx].cells[c_idx].text = cell_text

        # The body line spacing is tuned for prose; inside cells it just makes
        # rows tall and ragged.
        for row_cells in table.rows:
            for cell in row_cells.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_after = Pt(2)
                    para.paragraph_format.space_before = Pt(2)


def _md_to_docx(doc: Any, text: str, east_asian_font: str = "") -> None:
    """Convert markdown text to DOCX elements via HTML intermediate."""
    html = _md_to_html(text)
    renderer = _DocxMarkdownRenderer(doc, east_asian_font or docx_east_asian_font("en"))
    renderer.feed(html)
    renderer.close()


# ---------------------------------------------------------------------------
# DOCX renderer
# ---------------------------------------------------------------------------


def _render_docx(conv: Conversation, messages: list[Message], detail: DetailLevel, locale: str = "en", tz_name: str | None = None) -> bytes:
    """Render a conversation as a DOCX file and return the raw bytes."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise AppError(
            "docx_not_available",
            status_code=501,
            detail="DOCX export requires the python-docx package. "
            "Install with: uv pip install python-docx",
        )

    doc = Document()
    effective_mode = _resolve_effective_mode(conv, messages)

    _configure_docx_styles(doc, locale)
    east_asian = docx_east_asian_font(locale)

    title = _strip_emoji(conv.title or _t("untitled", locale))

    # Title
    doc.add_heading(title, level=1)

    # Metadata
    meta_text = f"{_t('mode', locale)}: {_mode_label(effective_mode, locale)}"
    meta_text += f"  |  {_t('created', locale)}: {_format_date(conv.created_at, tz_name)}"
    if detail == DetailLevel.FULL and conv.total_tokens:
        meta_text += f"  |  {_t('total_tokens', locale)}: {conv.total_tokens:,}"
    meta_para = doc.add_paragraph(meta_text)
    meta_para.runs[0].font.size = Pt(9)
    meta_para.runs[0].font.color.rgb = _hex_to_rgb(_MUTED)

    # Turns
    turns = _pair_messages(messages)
    for idx, turn in enumerate(turns, 1):
        doc.add_heading(f"{_t('turn', locale)} {idx}", level=2)

        user_msg: Message | None = turn.get("user")
        if user_msg:
            doc.add_paragraph(_t("user", locale), style=_DOCX_ROLE_STYLE)
            doc.add_paragraph(_strip_emoji(_strip_attachments(user_msg.content or "")))

        asst_msg: Message | None = turn.get("assistant")
        if asst_msg is None:
            continue

        doc.add_paragraph(_t("assistant", locale), style=_DOCX_ROLE_STYLE)

        if detail == DetailLevel.FULL:
            events = _extract_sse_events(asst_msg)
            done = _extract_done_event(events)

            if effective_mode == "dag":
                _render_docx_dag_details(doc, events, done, locale)
            else:
                _render_docx_react_details(doc, events, done, locale)

        # Final answer
        answer = asst_msg.content or ""
        if not answer and asst_msg.metadata_:
            answer = asst_msg.metadata_.get("answer", "")
        if answer:
            _md_to_docx(doc, answer, east_asian)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_monospace_paragraph(doc: Any, text: str, east_asian_font: str = "") -> None:
    """Add a paragraph with monospace font for code-like content."""
    try:
        from docx.shared import Pt
    except ImportError:
        doc.add_paragraph(text)
        return

    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    run.font.size = Pt(9)
    _apply_rpr_fonts(
        run._element.get_or_add_rPr(),
        DOCX_MONO_FONT,
        east_asian_font or docx_east_asian_font("en"),
    )


def _render_docx_react_details(
    doc: Any,
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
    locale: str = "en",
) -> None:
    """Add ReAct execution details to a DOCX document."""
    steps = _extract_react_steps(events)
    if not steps:
        return

    iterations = done.get("iterations", len(steps)) if done else len(steps)
    elapsed = done.get("elapsed", 0) if done else 0

    iter_label = _t("iteration" if iterations == 1 else "iterations", locale)
    doc.add_paragraph(
        f"{_t('execution_details', locale)}: {iterations} {iter_label}, {elapsed:.1f}s"
    )

    for i, step in enumerate(steps, 1):
        tool = _strip_emoji(step.get("tool_name", "unknown"))
        dur = step.get("iter_elapsed", 0)

        doc.add_paragraph(
            f"{_t('step', locale)} {i}: {tool} ({dur:.1f}s)", style="List Bullet"
        )

        reasoning = step.get("reasoning")
        if reasoning:
            doc.add_paragraph(f"{_t('reasoning', locale)}: {_strip_emoji(reasoning)}", style="List Bullet 2")

        args = step.get("tool_args")
        if args:
            _add_monospace_paragraph(doc, _strip_emoji(json.dumps(args, indent=2, ensure_ascii=False)))

        observation = step.get("observation")
        if observation:
            doc.add_paragraph(f"{_t('result', locale)}: {_strip_emoji(str(observation)[:500])}", style="List Bullet 2")


def _render_docx_dag_details(
    doc: Any,
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
    locale: str = "en",
) -> None:
    """Add DAG execution details to a DOCX document."""
    rounds = _extract_dag_rounds(events)

    for rnd in rounds:
        plan_steps = _extract_dag_plan(events, rnd)
        if plan_steps:
            doc.add_paragraph(f"{_t('plan', locale)} ({_t('round', locale)} {rnd}):")
            for ps in plan_steps:
                sid = _strip_emoji(ps.get("id", "?"))
                task = _strip_emoji(ps.get("task", ""))
                deps = ", ".join(ps.get("deps", [])) or _t("none", locale)
                doc.add_paragraph(
                    f"{sid}: {task} [{_t('dependencies', locale)}: {deps}]", style="List Bullet"
                )

    step_details = _extract_dag_step_details(events)
    for sid, info in step_details.items():
        task = _strip_emoji(info.get("task", ""))
        completed = info.get("completed")
        status = _strip_emoji(completed.get("status", "completed") if completed else "unknown")
        duration = completed.get("duration", 0) if completed else 0

        doc.add_paragraph(f"{sid}: {task} ({status}, {duration:.1f}s)")

        for it in info.get("iterations", []):
            it_num = it.get("iteration", "?")
            tool = _strip_emoji(it.get("tool_name", "unknown"))
            dur = it.get("iter_elapsed", 0)
            reasoning = _strip_emoji(it.get("reasoning", ""))
            doc.add_paragraph(
                f"{_t('iteration_label', locale)} {it_num}: {tool} ({dur:.1f}s) - {reasoning}",
                style="List Bullet",
            )

            observation = it.get("observation", "")
            if observation:
                doc.add_paragraph(
                    f"{_t('result', locale)}: {_strip_emoji(str(observation)[:500])}", style="List Bullet 2"
                )

        result = completed.get("result", "") if completed else ""
        if result:
            doc.add_paragraph(f"{_t('step_result', locale)}: {_strip_emoji(result)}", style="List Bullet")

    analysis = _extract_dag_analysis(events)
    if analysis:
        achieved = analysis.get("achieved", False)
        confidence = analysis.get("confidence", 0)
        reasoning = _strip_emoji(analysis.get("reasoning", ""))
        label = _t("goal_achieved", locale) if achieved else _t("goal_not_achieved", locale)
        doc.add_paragraph(f"{_t('analysis', locale)}: {label} ({confidence * 100:.0f}% {_t('confidence', locale)})")
        if reasoning:
            doc.add_paragraph(reasoning)


# ---------------------------------------------------------------------------
# PDF renderer
# ---------------------------------------------------------------------------


def _build_pdf_styles(fonts: PdfFonts) -> dict[str, Any]:
    """Build the ReportLab paragraph styles for the PDF export.

    Every style is created twice: ``key`` wraps character-by-character for CJK
    text, ``key_ltr`` keeps word-level wrapping for Latin-only text.  Callers
    pick between them with :func:`_pdf_style`.

    The size scale is deliberately five-tiered — document title, turn, role
    label, content headings, body — so headings written by the assistant stay
    visually subordinate to the transcript's own structure.
    """
    from reportlab.lib.enums import TA_LEFT  # type: ignore[import-untyped]
    from reportlab.lib.styles import ParagraphStyle  # type: ignore[import-untyped]
    from reportlab.lib.units import inch  # type: ignore[import-untyped]

    styles: dict[str, Any] = {}

    def add(key: str, **kwargs: Any) -> None:
        kwargs.setdefault("fontName", fonts.regular)
        kwargs.setdefault("alignment", TA_LEFT)
        kwargs.setdefault("textColor", _INK)
        styles[key] = ParagraphStyle(f"pdf_{key}", wordWrap="CJK", **kwargs)
        styles[f"{key}_ltr"] = ParagraphStyle(f"pdf_{key}_ltr", **kwargs)

    # --- document structure -------------------------------------------------
    add("title", fontName=fonts.bold, fontSize=20, leading=29, spaceAfter=6)
    add("meta", fontSize=8.5, leading=13, spaceAfter=4, textColor=_MUTED)
    add(
        "heading2",  # "Turn N"
        fontName=fonts.bold, fontSize=15, leading=22,
        spaceBefore=20, spaceAfter=9, textColor=_ACCENT,
    )
    add(
        "heading3",  # "User" / "Assistant"
        fontName=fonts.bold, fontSize=10, leading=15,
        spaceBefore=13, spaceAfter=5, textColor=_ACCENT,
    )

    # --- headings inside the assistant's markdown ---------------------------
    add("content_h1", fontName=fonts.bold, fontSize=13.5, leading=20, spaceBefore=14, spaceAfter=7)
    add("content_h2", fontName=fonts.bold, fontSize=12, leading=18, spaceBefore=12, spaceAfter=6)
    add("content_h3", fontName=fonts.bold, fontSize=11.5, leading=17, spaceBefore=10, spaceAfter=5)
    add("content_h4", fontName=fonts.bold, fontSize=11, leading=16, spaceBefore=9, spaceAfter=4)

    # --- body ---------------------------------------------------------------
    add("body", fontSize=10.5, leading=18, spaceAfter=8)
    add(
        "bullet", fontSize=10.5, leading=18, spaceAfter=4,
        leftIndent=0.32 * inch, bulletIndent=0.12 * inch,
    )
    add(
        "quote", fontSize=10, leading=17, spaceAfter=8,
        leftIndent=0.3 * inch, textColor=_MUTED,
    )
    # Sub-lines of a tool-call step, indented under their bullet.
    add(
        "detail", fontSize=9.5, leading=15, spaceAfter=4,
        leftIndent=0.62 * inch, textColor=_MUTED,
    )
    # Code blocks come in two indents: flush with body prose, and nested under a
    # tool-call bullet. Each needs a CJK variant because Courier has no CJK
    # coverage and would render Chinese source comments as blanks.
    for suffix, indent in (("", 0.25), ("_detail", 0.62)):
        add(
            f"code_cjk{suffix}", fontSize=9, leading=14, spaceAfter=8,
            leftIndent=indent * inch, backColor=_CODE_BG, textColor=_INK,
        )
        add(
            f"code{suffix}", fontName=fonts.mono, fontSize=8.5, leading=13, spaceAfter=8,
            leftIndent=indent * inch, backColor=_CODE_BG, textColor=_INK,
        )

    return styles


def _pdf_style(styles: dict[str, Any], key: str, text: str) -> Any:
    """Pick the CJK-wrapping or Latin-wrapping variant of a style."""
    if _has_cjk(text):
        return styles[key]
    return styles.get(f"{key}_ltr", styles[key])


class _PdfMarkdownRenderer(HTMLParser):
    """Parse HTML (converted from Markdown) and emit ReportLab flowables."""

    def __init__(self, styles: dict[str, Any], fonts: PdfFonts) -> None:
        super().__init__()
        self._styles = styles
        self._fonts = fonts
        self._font_name = fonts.regular
        self.flowables: list[Any] = []

        self._text_buf = ""
        self._bold = False
        self._italic = False
        self._code_inline = False
        self._in_pre = False
        self._pre_text = ""
        self._heading_level = 0
        self._heading_text = ""
        self._in_blockquote = False
        self._in_li = False
        self._list_stack: list[str] = []  # "ul" or "ol"
        self._ol_counters: list[int] = []
        self._in_table = False
        self._table_rows: list[list[str]] = []
        self._current_row: list[str] = []
        self._current_cell_text = ""
        self._in_td = False
        self._in_th = False
        self._href: str | None = None

    def _flush_text(self) -> None:
        """Flush accumulated text into a Paragraph flowable."""
        if not self._text_buf.strip():
            self._text_buf = ""
            return

        from reportlab.platypus import Paragraph  # type: ignore[import-untyped]

        key = "quote" if self._in_blockquote else "body"
        style = _pdf_style(self._styles, key, self._text_buf)
        self.flowables.append(Paragraph(self._text_buf, style))
        self._text_buf = ""

    def _wrap_inline(self, text: str) -> str:
        """Wrap text with inline formatting tags for ReportLab Paragraph."""
        # Escape XML entities
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if self._bold:
            text = f"<b>{text}</b>"
        if self._italic:
            text = f"<i>{text}</i>"
        if self._code_inline:
            # Courier has no CJK coverage — inline code containing Chinese has to
            # stay in the body face or it renders as blanks.
            face = self._fonts.mono if not _has_cjk(text) else self._fonts.regular
            text = f'<font name="{face}" size="9.5">{text}</font>'
        if self._href:
            text = f'<u><font color="#946B2D">{text}</font></u>'
        return text

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_dict = dict(attrs)

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._flush_text()
            self._heading_level = int(tag[1])
            self._heading_text = ""
        elif tag == "p":
            pass  # text accumulates in _text_buf
        elif tag in ("strong", "b"):
            self._bold = True
        elif tag in ("em", "i"):
            self._italic = True
        elif tag == "code":
            if not self._in_pre:
                self._code_inline = True
        elif tag == "pre":
            self._flush_text()
            self._in_pre = True
            self._pre_text = ""
        elif tag == "ul":
            self._flush_text()
            self._list_stack.append("ul")
        elif tag == "ol":
            self._flush_text()
            self._list_stack.append("ol")
            self._ol_counters.append(0)
        elif tag == "li":
            self._in_li = True
            self._text_buf = ""
        elif tag == "table":
            self._flush_text()
            self._in_table = True
            self._table_rows = []
        elif tag == "tr":
            self._current_row = []
        elif tag in ("td", "th"):
            self._in_td = tag == "td"
            self._in_th = tag == "th"
            self._current_cell_text = ""
        elif tag == "blockquote":
            self._flush_text()
            self._in_blockquote = True
        elif tag == "a":
            self._href = attrs_dict.get("href", "")
        elif tag == "hr":
            self._flush_text()
            from reportlab.platypus import HRFlowable

            self.flowables.append(HRFlowable(width="100%", thickness=0.5, color="#C8B898"))
        elif tag == "br":
            self._text_buf += "<br/>"

    def handle_endtag(self, tag: str) -> None:
        from reportlab.platypus import Paragraph, Preformatted, Spacer

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            # Headings written by the assistant sit *below* the transcript's own
            # turn/role headings, so h1 maps to the first content tier rather
            # than competing with the document title.
            style_key = f"content_h{min(self._heading_level, 4)}"
            # Escape XML entities in heading text
            safe = self._heading_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            self.flowables.append(Paragraph(safe, _pdf_style(self._styles, style_key, safe)))
            self._heading_level = 0
            self._heading_text = ""
        elif tag == "p":
            self._flush_text()
        elif tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False
        elif tag == "code":
            if not self._in_pre:
                self._code_inline = False
        elif tag == "pre":
            self._in_pre = False
            safe = self._pre_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            code_key = "code_cjk" if _has_cjk(safe) else "code"
            self.flowables.append(Preformatted(safe, self._styles[code_key]))
            self._pre_text = ""
        elif tag == "ul":
            if self._list_stack:
                self._list_stack.pop()
        elif tag == "ol":
            if self._list_stack:
                self._list_stack.pop()
            if self._ol_counters:
                self._ol_counters.pop()
        elif tag == "li":
            self._in_li = False
            text = self._text_buf.strip()
            self._text_buf = ""
            if text:
                prefix = "\u2022 "
                if self._list_stack and self._list_stack[-1] == "ol":
                    if self._ol_counters:
                        self._ol_counters[-1] += 1
                        prefix = f"{self._ol_counters[-1]}. "
                self.flowables.append(
                    Paragraph(f"{prefix}{text}", _pdf_style(self._styles, "bullet", text))
                )
        elif tag in ("td", "th"):
            self._current_row.append(self._current_cell_text.strip())
            self._in_td = False
            self._in_th = False
            self._current_cell_text = ""
        elif tag == "tr":
            self._table_rows.append(self._current_row)
            self._current_row = []
        elif tag == "table":
            self._in_table = False
            if self._table_rows:
                self._emit_table()
        elif tag == "blockquote":
            self._flush_text()
            self._in_blockquote = False
        elif tag == "a":
            self._href = None

    def handle_charref(self, name: str) -> None:
        """Convert HTML numeric character references to characters and route
        through ``handle_data`` so emoji stripping applies uniformly."""
        try:
            codepoint = int(name[1:], 16) if name.startswith(("x", "X")) else int(name)
            self.handle_data(chr(codepoint))
        except (ValueError, OverflowError):
            pass

    def handle_data(self, data: str) -> None:
        data = _strip_emoji(data)
        if self._heading_level:
            self._heading_text += data
        elif self._in_pre:
            self._pre_text += data
        elif self._in_td or self._in_th:
            self._current_cell_text += data
        elif self._in_table:
            pass
        elif self._in_li:
            self._text_buf += self._wrap_inline(data)
        else:
            self._text_buf += self._wrap_inline(data)

    def _emit_table(self) -> None:
        from reportlab.lib import colors  # type: ignore[import-untyped]
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, Table, TableStyle

        if not self._table_rows:
            return

        n_cols = max(len(r) for r in self._table_rows)

        # Calculate available width (A4 width minus margins)
        available_width = 6.5 * inch  # A4 width ~8.27in - 2*0.75in margins
        col_width = available_width / n_cols
        col_widths = [col_width] * n_cols

        table_data: list[list[Any]] = []
        for r_idx, row in enumerate(self._table_rows):
            cells = []
            for i in range(n_cols):
                cell_text = row[i] if i < len(row) else ""
                safe = cell_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                # Bold header row
                if r_idx == 0:
                    safe = f"<b>{safe}</b>"
                cells.append(Paragraph(safe, _pdf_style(self._styles, "body", safe)))
            table_data.append(cells)

        if not table_data:
            return

        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.78, 0.72, 0.60)),
            ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.96, 0.93, 0.87)),
            ("FONTNAME", (0, 0), (-1, -1), self._fonts.regular),
            ("FONTNAME", (0, 0), (-1, 0), self._fonts.bold),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("LEADING", (0, 0), (-1, -1), 14),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ]))
        self.flowables.append(table)


def _md_to_pdf_flowables(text: str, styles: dict[str, Any], fonts: PdfFonts) -> list[Any]:
    """Convert markdown text to a list of ReportLab flowables via HTML intermediate."""
    html = _md_to_html(text)
    renderer = _PdfMarkdownRenderer(styles, fonts)
    renderer.feed(html)
    renderer.close()
    # Flush any remaining text
    renderer._flush_text()
    return renderer.flowables


def _render_pdf(conv: Conversation, messages: list[Message], detail: DetailLevel, locale: str = "en", tz_name: str | None = None) -> bytes:
    """Render a conversation as a PDF file and return the raw bytes."""
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore[import-untyped]
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer
        from reportlab.platypus import HRFlowable
    except ImportError:
        raise AppError(
            "pdf_not_available",
            status_code=501,
            detail="PDF export requires the reportlab package. "
            "Install with: uv pip install reportlab",
        )

    fonts = resolve_pdf_fonts()
    styles = _build_pdf_styles(fonts)
    effective_mode = _resolve_effective_mode(conv, messages)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    flowables: list[Any] = []
    title = _strip_emoji(conv.title or _t("untitled", locale))

    # Title
    safe_title = title.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    flowables.append(Paragraph(safe_title, _pdf_style(styles, "title", safe_title)))

    # Metadata
    meta_text = f"{_t('mode', locale)}: {_mode_label(effective_mode, locale)}"
    meta_text += f"  |  {_t('created', locale)}: {_format_date(conv.created_at, tz_name)}"
    if detail == DetailLevel.FULL and conv.total_tokens:
        meta_text += f"  |  {_t('total_tokens', locale)}: {conv.total_tokens:,}"
    safe_meta = meta_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    flowables.append(Paragraph(safe_meta, _pdf_style(styles, "meta", safe_meta)))
    flowables.append(Spacer(1, 10))
    flowables.append(HRFlowable(width="100%", thickness=1, color=_ACCENT))

    # Turns
    turns = _pair_messages(messages)
    for idx, turn in enumerate(turns, 1):
        turn_label = f"{_t('turn', locale)} {idx}"
        flowables.append(Paragraph(turn_label, _pdf_style(styles, "heading2", turn_label)))

        user_msg: Message | None = turn.get("user")
        if user_msg:
            role = _t("user", locale)
            flowables.append(Paragraph(role, _pdf_style(styles, "heading3", role)))
            user_text = _strip_emoji(_strip_attachments(user_msg.content or ""))
            safe_user = user_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            flowables.append(Paragraph(safe_user, _pdf_style(styles, "body", safe_user)))

        asst_msg: Message | None = turn.get("assistant")
        if asst_msg is None:
            flowables.append(Spacer(1, 6))
            flowables.append(HRFlowable(width="100%", thickness=0.5, color=_RULE))
            continue

        role = _t("assistant", locale)
        flowables.append(Paragraph(role, _pdf_style(styles, "heading3", role)))

        if detail == DetailLevel.FULL:
            events = _extract_sse_events(asst_msg)
            done = _extract_done_event(events)

            if effective_mode == "dag":
                _render_pdf_dag_details(flowables, events, done, styles, fonts, locale)
            else:
                _render_pdf_react_details(flowables, events, done, styles, fonts, locale)

        # Final answer
        answer = asst_msg.content or ""
        if not answer and asst_msg.metadata_:
            answer = asst_msg.metadata_.get("answer", "")
        if answer:
            answer_flowables = _md_to_pdf_flowables(answer, styles, fonts)
            flowables.extend(answer_flowables)

        flowables.append(Spacer(1, 8))
        flowables.append(HRFlowable(width="100%", thickness=0.5, color=_RULE))

    # Guard against empty document (ReportLab raises on no flowables)
    if not flowables:
        empty = _t("empty_conversation", locale)
        flowables.append(Paragraph(empty, _pdf_style(styles, "body", empty)))

    doc.build(flowables)
    return buf.getvalue()


def _render_pdf_react_details(
    flowables: list[Any],
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
    styles: dict[str, Any],
    fonts: PdfFonts,
    locale: str = "en",
) -> None:
    """Add ReAct execution details as PDF flowables."""
    from reportlab.platypus import Paragraph, Preformatted, Spacer

    steps = _extract_react_steps(events)
    if not steps:
        return

    iterations = done.get("iterations", len(steps)) if done else len(steps)
    elapsed = done.get("elapsed", 0) if done else 0

    iter_label = _t("iteration" if iterations == 1 else "iterations", locale)
    summary = f"<b>{_t('execution_details', locale)}:</b> {iterations} {iter_label}, {elapsed:.1f}s"
    flowables.append(Paragraph(summary, _pdf_style(styles, "meta", summary)))

    for i, step in enumerate(steps, 1):
        tool = _strip_emoji(step.get("tool_name", "unknown"))
        dur = step.get("iter_elapsed", 0)
        safe_tool = tool.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        line = f"\u2022 {_t('step', locale)} {i}: <b>{safe_tool}</b> ({dur:.1f}s)"
        flowables.append(Paragraph(line, _pdf_style(styles, "bullet", line)))

        reasoning = step.get("reasoning")
        if reasoning:
            safe_r = _strip_emoji(str(reasoning)).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            line = f"{_t('reasoning', locale)}: {safe_r}"
            flowables.append(Paragraph(line, _pdf_style(styles, "detail", line)))

        args = step.get("tool_args")
        if args:
            args_text = _strip_emoji(json.dumps(args, indent=2, ensure_ascii=False))
            safe_args = args_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            code_key = "code_cjk_detail" if _has_cjk(safe_args) else "code_detail"
            flowables.append(Preformatted(safe_args, styles[code_key]))

        observation = step.get("observation")
        if observation:
            obs_text = _strip_emoji(str(observation)[:500])
            safe_obs = obs_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            line = f"{_t('result', locale)}: {safe_obs}"
            flowables.append(Paragraph(line, _pdf_style(styles, "detail", line)))

    flowables.append(Spacer(1, 8))


def _render_pdf_dag_details(
    flowables: list[Any],
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
    styles: dict[str, Any],
    fonts: PdfFonts,
    locale: str = "en",
) -> None:
    """Add DAG execution details as PDF flowables."""
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

    rounds = _extract_dag_rounds(events)

    for rnd in rounds:
        plan_steps = _extract_dag_plan(events, rnd)
        if plan_steps:
            plan_label = f"<b>{_t('plan', locale)} ({_t('round', locale)} {rnd}):</b>"
            flowables.append(Paragraph(
                plan_label, _pdf_style(styles, "meta", plan_label)
            ))

            # Build table: Step | Task | Dependencies
            header = [
                Paragraph(f"<b>{_t('step', locale)}</b>", styles["body"]),
                Paragraph(f"<b>{_t('task', locale)}</b>", styles["body"]),
                Paragraph(f"<b>{_t('dependencies', locale)}</b>", styles["body"]),
            ]
            table_data = [header]
            for ps in plan_steps:
                sid = _strip_emoji(ps.get("id", "?"))
                task = _strip_emoji(ps.get("task", ""))
                deps = ", ".join(ps.get("deps", [])) or "\u2014"
                safe_sid = str(sid).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                safe_task = task.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                safe_deps = deps.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                table_data.append([
                    Paragraph(safe_sid, _pdf_style(styles, "body", safe_sid)),
                    Paragraph(safe_task, _pdf_style(styles, "body", safe_task)),
                    Paragraph(safe_deps, _pdf_style(styles, "body", safe_deps)),
                ])

            # Weighted column widths: Step narrow, Task wide, Dependencies medium
            available_width = 6.5 * inch
            col_widths = [available_width * 0.12, available_width * 0.60, available_width * 0.28]

            table = Table(table_data, colWidths=col_widths, repeatRows=1)
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.78, 0.72, 0.60)),
                ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.96, 0.93, 0.87)),
                ("FONTNAME", (0, 0), (-1, -1), fonts.regular),
                ("FONTNAME", (0, 0), (-1, 0), fonts.bold),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ]))
            flowables.append(table)
            flowables.append(Spacer(1, 8))

    step_details = _extract_dag_step_details(events)
    for sid, info in step_details.items():
        task = _strip_emoji(info.get("task", ""))
        completed = info.get("completed")
        status = _strip_emoji(completed.get("status", "completed") if completed else "unknown")
        duration = completed.get("duration", 0) if completed else 0
        safe_sid = _strip_emoji(str(sid)).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        safe_task = task.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        safe_status = status.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        line = f"<b>{safe_sid}:</b> {safe_task} ({safe_status}, {duration:.1f}s)"
        flowables.append(Paragraph(line, _pdf_style(styles, "body", line)))

        for it in info.get("iterations", []):
            it_num = it.get("iteration", "?")
            tool = _strip_emoji(it.get("tool_name", "unknown"))
            dur = it.get("iter_elapsed", 0)
            reasoning = _strip_emoji(it.get("reasoning", ""))
            observation = it.get("observation", "")
            safe_tool = str(tool).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            safe_reason = str(reasoning).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

            line = (
                f"\u2022 {_t('iteration_label', locale)} {it_num}: "
                f"<b>{safe_tool}</b> ({dur:.1f}s) \u2014 {safe_reason}"
            )
            flowables.append(Paragraph(line, _pdf_style(styles, "bullet", line)))

            if observation:
                obs_text = _strip_emoji(str(observation)[:500])
                safe_obs = obs_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                obs_line = f"{_t('result', locale)}: {safe_obs}"
                flowables.append(Paragraph(obs_line, _pdf_style(styles, "detail", obs_line)))

        result = completed.get("result", "") if completed else ""
        if result:
            safe_result = _strip_emoji(str(result)).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            res_line = f"{_t('step_result', locale)}: {safe_result}"
            flowables.append(Paragraph(res_line, _pdf_style(styles, "bullet", res_line)))
        flowables.append(Spacer(1, 4))

    analysis = _extract_dag_analysis(events)
    if analysis:
        achieved = analysis.get("achieved", False)
        confidence = analysis.get("confidence", 0)
        reasoning = _strip_emoji(analysis.get("reasoning", ""))
        label = _t("goal_achieved", locale) if achieved else _t("goal_not_achieved", locale)
        line = f"<b>{_t('analysis', locale)}:</b> {label} ({confidence * 100:.0f}% {_t('confidence', locale)})"
        flowables.append(Paragraph(line, _pdf_style(styles, "meta", line)))
        if reasoning:
            safe_reason = str(reasoning).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            flowables.append(Paragraph(safe_reason, _pdf_style(styles, "quote", safe_reason)))
        flowables.append(Spacer(1, 8))


# ---------------------------------------------------------------------------
# Endpoint
# ---------------------------------------------------------------------------

_CONTENT_TYPES: dict[ExportFormat, str] = {
    ExportFormat.MD: "text/markdown; charset=utf-8",
    ExportFormat.TXT: "text/plain; charset=utf-8",
    ExportFormat.DOCX: (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    ),
    ExportFormat.PDF: "application/pdf",
}


@router.get("/{conversation_id}/export")
async def export_conversation(
    conversation_id: str,
    format: ExportFormat = Query(..., description="Export format: md, txt, docx, or pdf"),
    detail: DetailLevel = Query(
        DetailLevel.FULL, description="Detail level: full or summary"
    ),
    locale: str = Query("en", description="Locale for labels: en or zh"),
    current_user: User = Depends(get_current_user),  # noqa: B008
    db: AsyncSession = Depends(get_session),  # noqa: B008
) -> StreamingResponse:
    """Export a conversation as a downloadable file.

    Supports Markdown, plain text, DOCX, and PDF formats.  The ``detail``
    parameter controls whether tool execution details are included
    (``full``) or only the final answers (``summary``).
    """
    # Fetch conversation with messages, verify ownership
    result = await db.execute(
        select(Conversation)
        .where(
            Conversation.id == conversation_id,
            Conversation.user_id == current_user.id,
        )
        .options(selectinload(Conversation.messages))
    )
    conv = result.scalar_one_or_none()
    if conv is None:
        raise AppError("conversation_not_found", status_code=404)

    # Sort messages chronologically
    messages = sorted(conv.messages, key=lambda m: m.created_at)

    # Build filename
    tz_name: str | None = getattr(current_user, "timezone", None)
    safe_title = _sanitize_filename(conv.title or "conversation")
    date_str = _format_date_compact(conv.created_at, tz_name)
    ext = format.value
    detail_suffix = "_full" if detail == DetailLevel.FULL else ""
    filename = f"{safe_title}_{date_str}{detail_suffix}.{ext}"

    # Render content
    if format == ExportFormat.PDF:
        content_bytes = _render_pdf(conv, messages, detail, locale, tz_name)
        stream = io.BytesIO(content_bytes)
    elif format == ExportFormat.DOCX:
        content_bytes = _render_docx(conv, messages, detail, locale, tz_name)
        stream = io.BytesIO(content_bytes)
    elif format == ExportFormat.TXT:
        text = _render_txt(conv, messages, detail, locale, tz_name)
        stream = io.BytesIO(text.encode("utf-8"))
    else:  # MD
        text = _render_md(conv, messages, detail, locale, tz_name)
        stream = io.BytesIO(text.encode("utf-8"))

    content_type = _CONTENT_TYPES[format]

    # RFC 5987: use ASCII fallback + UTF-8 encoded filename for non-ASCII titles
    ascii_filename = filename.encode("ascii", errors="ignore").decode("ascii") or f"export.{ext}"
    utf8_filename = quote(filename)
    disposition = f"attachment; filename=\"{ascii_filename}\"; filename*=UTF-8''{utf8_filename}"

    return StreamingResponse(
        stream,
        media_type=content_type,
        headers={
            "Content-Disposition": disposition,
        },
    )
