"""DOCX document loader using python-docx."""

from __future__ import annotations

import asyncio
from pathlib import Path

from .base import BaseLoader, LoadedDocument


class DOCXLoader(BaseLoader):
    """Load DOCX files using python-docx."""

    async def load(self, path: Path) -> list[LoadedDocument]:
        return await asyncio.to_thread(self._load_sync, path)

    def _load_sync(self, path: Path) -> list[LoadedDocument]:
        import docx

        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        if not content.strip():
            return []

        return [
            LoadedDocument(
                content=content,
                metadata={"source": str(path)},
            )
        ]
